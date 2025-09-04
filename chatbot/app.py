import streamlit as st
import os
import json
from typing import List, Dict, Any, Optional
from dataclasses import dataclass
import numpy as np
from sentence_transformers import SentenceTransformer
import faiss
import requests
from groq import Groq
import tempfile
import PyPDF2
import docx
from io import BytesIO
import re

@dataclass
class Document:
    content: str
    metadata: Dict[str, Any] = None

class ConstructionRAGSystem:
    def __init__(self, groq_api_key: str, embedding_model: str = "all-MiniLM-L6-v2"):
        """Initialize Construction-focused RAG system with Groq API and embedding model"""
        self.groq_client = Groq(api_key=groq_api_key)
        
        # Initialize embedding model with caching
        @st.cache_resource
        def load_embedding_model():
            return SentenceTransformer(embedding_model)
        
        self.embedding_model = load_embedding_model()
        
        # Initialize vector database (FAISS)
        self.dimension = self.embedding_model.get_sentence_embedding_dimension()
        self.index = faiss.IndexFlatIP(self.dimension)
        
        # Storage for documents
        self.documents: List[Document] = []
        
        # Construction-related keywords for filtering
        self.construction_keywords = [
            'construction', 'building', 'concrete', 'steel', 'foundation', 'excavation',
            'architecture', 'engineering', 'contractor', 'blueprint', 'scaffold',
            'cement', 'rebar', 'structural', 'hvac', 'plumbing', 'electrical',
            'roofing', 'flooring', 'framing', 'insulation', 'drywall', 'masonry',
            'crane', 'bulldozer', 'excavator', 'project management', 'site safety',
            'building code', 'permit', 'inspection', 'renovation', 'demolition',
            'civil engineering', 'surveying', 'geotechnical', 'soil', 'aggregate',
            'asphalt', 'paving', 'bridge', 'tunnel', 'highway', 'infrastructure',
            'residential', 'commercial', 'industrial', 'apartment', 'office building',
            'warehouse', 'factory', 'hospital', 'school', 'mall', 'parking garage'
        ]
    
    def is_construction_related(self, query: str) -> bool:
        """Check if the query is construction-related"""
        query_lower = query.lower()
        
        # Check for construction keywords
        for keyword in self.construction_keywords:
            if keyword in query_lower:
                return True
        
        # Check for construction-related patterns
        construction_patterns = [
            r'\b(build|construct|design|plan)\b',
            r'\b(material|tool|equipment)\b.*\b(construction|building)\b',
            r'\b(how to|what is|explain)\b.*\b(construction|building|concrete|steel)\b',
            r'\b(cost|price|budget)\b.*\b(construction|building)\b',
            r'\b(safety|regulation|code)\b.*\b(construction|building)\b'
        ]
        
        for pattern in construction_patterns:
            if re.search(pattern, query_lower):
                return True
        
        return False
    
    def chunk_text(self, text: str, chunk_size: int = 500, overlap: int = 50) -> List[str]:
        """Split text into overlapping chunks"""
        words = text.split()
        chunks = []
        
        for i in range(0, len(words), chunk_size - overlap):
            chunk = " ".join(words[i:i + chunk_size])
            if chunk.strip():
                chunks.append(chunk)
                
        return chunks
    
    def add_documents(self, texts: List[str], metadatas: List[Dict] = None):
        """Add documents to the vector database"""
        if metadatas is None:
            metadatas = [{}] * len(texts)
            
        # Create document objects
        docs = [Document(content=text, metadata=meta) 
                for text, meta in zip(texts, metadatas)]
        
        # Generate embeddings
        embeddings = self.embedding_model.encode(texts)
        
        # Normalize embeddings for cosine similarity
        embeddings = embeddings / np.linalg.norm(embeddings, axis=1, keepdims=True)
        
        # Add to FAISS index
        self.index.add(embeddings.astype('float32'))
        
        # Store documents
        self.documents.extend(docs)
        
        return len(texts)
    
    def retrieve_relevant_chunks(self, query: str, top_k: int = 3) -> List[Dict]:
        """Retrieve top-k relevant document chunks if documents exist"""
        if self.index.ntotal == 0:
            return []
            
        # Generate query embedding
        query_embedding = self.embedding_model.encode([query])
        query_embedding = query_embedding / np.linalg.norm(query_embedding, axis=1, keepdims=True)
        
        # Search in vector database
        scores, indices = self.index.search(query_embedding.astype('float32'), top_k)
        
        # Retrieve relevant documents with scores
        relevant_chunks = []
        for score, idx in zip(scores[0], indices[0]):
            if idx < len(self.documents) and score > 0.2:  # Lower threshold for reference
                relevant_chunks.append({
                    'content': self.documents[idx].content,
                    'metadata': self.documents[idx].metadata,
                    'score': float(score)
                })
        
        return relevant_chunks
    
    def build_construction_prompt(self, query: str, relevant_chunks: List[Dict], chat_history: List[Dict]) -> str:
        """Build specialized construction prompt"""
        
        # System prompt for construction expertise
        system_prompt = """You are a construction industry expert and consultant with deep knowledge in:
- Construction methods, materials, and techniques
- Building codes, regulations, and safety standards  
- Project management and cost estimation
- Structural engineering and design principles
- Construction equipment and tools
- Site safety and risk management
- Quality control and inspection procedures
- Sustainable construction practices

You provide accurate, practical, and professional advice for construction-related questions.
Answer based on industry best practices and current standards."""
        
        # Add conversation history
        history_text = ""
        if chat_history:
            history_text = "\n\nPrevious conversation:\n"
            for exchange in chat_history[-2:]:  # Last 2 exchanges
                history_text += f"Human: {exchange['human']}\nAssistant: {exchange['assistant']}\n"
        
        # Add document context as reference (not primary source)
        context_text = ""
        if relevant_chunks:
            context_text = "\n\nAdditional reference materials (use as supporting information):\n"
            for i, chunk in enumerate(relevant_chunks, 1):
                source_info = chunk['metadata'].get('source', 'Document')
                context_text += f"[Reference {i} from {source_info}]: {chunk['content']}\n\n"
            context_text += "\nNote: Use the above references to supplement your expert knowledge, but provide comprehensive answers based on your construction expertise.\n"
        
        # Build final prompt
        prompt = f"""{system_prompt}
        {history_text}
        {context_text}
        
Human: {query}
Assistant:"""
        
        return prompt
    
    def generate_construction_response(self, query: str, chat_history: List[Dict] = None) -> tuple[str, List[Dict], bool]:
        """Generate response for construction queries"""
        if chat_history is None:
            chat_history = []
        
        # Check if query is construction-related
        if not self.is_construction_related(query):
            return ("I'm specialized in construction-related questions only. Please ask about construction methods, materials, building techniques, project management, safety, codes, or any other construction industry topics.", [], False)
        
        # Retrieve relevant chunks (if any documents exist)
        relevant_chunks = self.retrieve_relevant_chunks(query, top_k=3)
        
        # Build prompt
        prompt = self.build_construction_prompt(query, relevant_chunks, chat_history)
        
        try:
            # Generate response using Groq
            response = self.groq_client.chat.completions.create(
                model="llama-3.3-70b-versatile",

                messages=[
                    {"role": "user", "content": prompt}
                ],
                temperature=0.7,
                max_tokens=1500
            )
            
            answer = response.choices[0].message.content
            
            return answer, relevant_chunks, True
            
        except Exception as e:
            return f"Error generating response: {str(e)}", [], False

def extract_text_from_pdf(pdf_file):
    """Extract text from PDF file"""
    try:
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text
    except Exception as e:
        st.error(f"Error reading PDF: {str(e)}")
        return ""

def extract_text_from_docx(docx_file):
    """Extract text from DOCX file"""
    try:
        doc = docx.Document(docx_file)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    except Exception as e:
        st.error(f"Error reading DOCX: {str(e)}")
        return ""

def main():
    st.set_page_config(
        page_title="Construction AI Assistant",
        page_icon="🏗️",
        layout="wide"
    )
    
    st.title("🏗️ Construction Industry AI Assistant")
    st.markdown("**Specialized AI assistant for construction-related questions with optional document reference**")
    
    # Information banner
    st.info("""
    🎯 **How this works:**
    - Ask any construction-related question and get expert answers
    - Optionally upload documents for additional reference context
    - Documents serve as supporting material, not the primary knowledge source
    - Only construction industry questions are accepted
    """)
    
    # Sidebar for configuration
    with st.sidebar:
        st.header("⚙️ Configuration")
        
        # Groq API Key input
        groq_api_key = st.text_input(
            "Groq API Key", 
            type="password",
            help="Get your API key from https://console.groq.com/keys"
        )
        
        if not groq_api_key:
            st.warning("Please enter your Groq API key to continue.")
            st.stop()
        
        st.header("📄 Optional Document Upload")
        st.markdown("Upload construction-related documents for additional reference:")
        
        # File uploader
        uploaded_files = st.file_uploader(
            "Upload reference documents",
            type=['txt', 'pdf', 'docx'],
            accept_multiple_files=True,
            help="Upload construction manuals, codes, specifications, etc."
        )
        
        # Parameters
        st.header("🔧 Parameters")
        chunk_size = st.slider("Chunk Size", 300, 800, 500)
        overlap = st.slider("Overlap", 25, 100, 50)
        
        st.header("🏗️ Construction Topics")
        st.markdown("""
        **Supported areas:**
        - Building methods & materials
        - Structural engineering
        - Project management
        - Safety & regulations  
        - Cost estimation
        - Construction equipment
        - Building codes
        - Site management
        - Quality control
        """)
    
    # Initialize RAG system
    if 'construction_rag' not in st.session_state:
        try:
            st.session_state.construction_rag = ConstructionRAGSystem(groq_api_key)
        except Exception as e:
            st.error(f"Error initializing system: {str(e)}")
            st.stop()
    
    # Initialize chat history
    if 'chat_history' not in st.session_state:
        st.session_state.chat_history = []
    
    # Process uploaded files
    if uploaded_files:
        with st.expander("📋 Document Processing", expanded=True):
            total_chunks = 0
            
            for uploaded_file in uploaded_files:
                st.write(f"Processing: {uploaded_file.name}")
                
                # Extract text based on file type
                if uploaded_file.type == "text/plain":
                    text = str(uploaded_file.read(), "utf-8")
                elif uploaded_file.type == "application/pdf":
                    text = extract_text_from_pdf(uploaded_file)
                elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                    text = extract_text_from_docx(uploaded_file)
                else:
                    st.error(f"Unsupported file type: {uploaded_file.type}")
                    continue
                
                if text.strip():
                    # Chunk the text
                    chunks = st.session_state.construction_rag.chunk_text(
                        text, chunk_size=chunk_size, overlap=overlap
                    )
                    
                    # Create metadata
                    metadatas = [{"source": uploaded_file.name, "chunk_id": i} 
                               for i in range(len(chunks))]
                    
                    # Add to vector database
                    num_added = st.session_state.construction_rag.add_documents(chunks, metadatas)
                    total_chunks += num_added
                    
                    st.success(f"✅ Added {num_added} reference chunks from {uploaded_file.name}")
                else:
                    st.warning(f"⚠️ No text extracted from {uploaded_file.name}")
            
            if total_chunks > 0:
                st.info(f"📊 Total reference chunks available: {total_chunks}")
    
    # Chat interface
    st.header("💬 Ask Construction Questions")
    
    # Sample questions
    with st.expander("💡 Example Questions"):
        example_questions = [
            "What are the standard concrete mix ratios for different strength requirements?",
            "How do you ensure safety when working at height on construction sites?",
            "What are the key factors in construction project cost estimation?",
            "Explain the difference between various foundation types and when to use each",
            "What are the latest building codes for seismic design?",
            "How to properly cure concrete in different weather conditions?",
            "What equipment is needed for a typical residential construction project?"
        ]
        
        for question in example_questions:
            if st.button(question, key=f"example_{hash(question)}"):
                st.session_state.example_question = question
    
    # Display chat history
    for exchange in st.session_state.chat_history:
        with st.chat_message("user"):
            st.write(exchange["human"])
        
        with st.chat_message("assistant"):
            st.write(exchange["assistant"])
            
            # Show references if available
            if exchange.get("sources"):
                with st.expander("📚 Reference Materials Used"):
                    for i, source in enumerate(exchange["sources"], 1):
                        st.write(f"**Reference {i}** (Relevance: {source['score']:.3f})")
                        st.write(f"From: {source['metadata'].get('source', 'Unknown')}")
                        st.write(source['content'][:300] + "..." if len(source['content']) > 300 else source['content'])
                        st.write("---")
    
    # Handle example question
    if 'example_question' in st.session_state:
        prompt = st.session_state.example_question
        del st.session_state.example_question
        
        # Process the example question
        with st.chat_message("user"):
            st.write(prompt)
        
        with st.chat_message("assistant"):
            with st.spinner("Generating construction expertise..."):
                response, sources, is_valid = st.session_state.construction_rag.generate_construction_response(
                    prompt, st.session_state.chat_history
                )
            
            st.write(response)
            
            # Show references if available
            if sources and is_valid:
                with st.expander("📚 Reference Materials Used"):
                    for i, source in enumerate(sources, 1):
                        st.write(f"**Reference {i}** (Relevance: {source['score']:.3f})")
                        st.write(f"From: {source['metadata'].get('source', 'Unknown')}")
                        st.write(source['content'][:300] + "..." if len(source['content']) > 300 else source['content'])
                        st.write("---")
        
        # Add to chat history
        st.session_state.chat_history.append({
            "human": prompt,
            "assistant": response,
            "sources": sources if is_valid else []
        })
    
    # Chat input
    if prompt := st.chat_input("Ask your construction question here..."):
        # Display user message
        with st.chat_message("user"):
            st.write(prompt)
        
        # Generate and display assistant response
        with st.chat_message("assistant"):
            with st.spinner("Generating construction expertise..."):
                response, sources, is_valid = st.session_state.construction_rag.generate_construction_response(
                    prompt, st.session_state.chat_history
                )
            
            st.write(response)
            
            # Show references if available
            if sources and is_valid:
                with st.expander("📚 Reference Materials Used"):
                    for i, source in enumerate(sources, 1):
                        st.write(f"**Reference {i}** (Relevance: {source['score']:.3f})")
                        st.write(f"From: {source['metadata'].get('source', 'Unknown')}")
                        st.write(source['content'][:300] + "..." if len(source['content']) > 300 else source['content'])
                        st.write("---")
        
        # Add to chat history
        st.session_state.chat_history.append({
            "human": prompt,
            "assistant": response,
            "sources": sources if is_valid else []
        })
        
        # Keep only last 8 exchanges
        if len(st.session_state.chat_history) > 8:
            st.session_state.chat_history = st.session_state.chat_history[-8:]
    
    # Clear chat button
    if st.sidebar.button("🗑️ Clear Chat History"):
        st.session_state.chat_history = []
        st.rerun()
    
    # Footer
    st.markdown("---")
    st.markdown("🏗️ **Construction AI Assistant** - Powered by Mistral via Groq | Specialized for Construction Industry")

if __name__ == "__main__":
    main()